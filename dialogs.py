# dialogs.py
from PyQt6.QtWidgets import (QDialog, QVBoxLayout, QFormLayout, QLabel,
                             QLineEdit, QSpinBox, QDoubleSpinBox, QPushButton,
                             QDialogButtonBox, QTableWidget, QTableWidgetItem,
                             QHBoxLayout, QTextEdit, QMessageBox, QFileDialog)
from PyQt6.QtCore import Qt
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
from datetime import datetime
import os
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import COMPANY_NAME, DEVELOPER_NAME, DEVELOPER_WHATSAPP, COMPANY_LOGO_PATH

class AddProductDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add Product")
        self.setModal(True)
        self.setFixedSize(500, 400)
        
        layout = QVBoxLayout()
        
        form_layout = QFormLayout()
        
        self.product_name = QLineEdit()
        self.product_name.setPlaceholderText("Enter product name")
        form_layout.addRow("Product Name:", self.product_name)
        
        self.purchase_price = QDoubleSpinBox()
        self.purchase_price.setRange(0, 999999)
        self.purchase_price.setPrefix("$ ")
        form_layout.addRow("Purchase Price:", self.purchase_price)
        
        self.selling_price = QDoubleSpinBox()
        self.selling_price.setRange(0, 999999)
        self.selling_price.setPrefix("$ ")
        form_layout.addRow("Selling Price:", self.selling_price)
        
        self.quantity = QSpinBox()
        self.quantity.setRange(0, 999999)
        form_layout.addRow("Quantity:", self.quantity)
        
        self.profit_label = QLabel("$ 0.00")
        self.profit_label.setStyleSheet("font-weight: bold; color: #4CAF50;")
        form_layout.addRow("Total Profit:", self.profit_label)
        
        layout.addLayout(form_layout)
        
        # Connect signals for auto profit calculation
        self.purchase_price.valueChanged.connect(self.calculate_profit)
        self.selling_price.valueChanged.connect(self.calculate_profit)
        self.quantity.valueChanged.connect(self.calculate_profit)
        
        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
        self.setLayout(layout)
    
    def calculate_profit(self):
        purchase = self.purchase_price.value()
        selling = self.selling_price.value()
        quantity = self.quantity.value()
        profit = (selling - purchase) * quantity
        self.profit_label.setText(f"$ {profit:.2f}")
    
    def get_product_data(self):
        return {
            'name': self.product_name.text(),
            'purchase_price': self.purchase_price.value(),
            'selling_price': self.selling_price.value(),
            'quantity': self.quantity.value(),
            'profit': (self.selling_price.value() - self.purchase_price.value()) * self.quantity.value()
        }

class InvoiceDialog(QDialog):
    def __init__(self, parent=None, invoice_data=None):
        super().__init__(parent)
        self.parent = parent
        self.invoice_data = invoice_data
        self.setWindowTitle(f"Invoice #{invoice_data['invoice_no']}")
        self.setModal(True)
        self.setMinimumSize(800, 600)
        
        layout = QVBoxLayout()
        
        # Invoice display area
        self.invoice_text = QTextEdit()
        self.invoice_text.setReadOnly(True)
        layout.addWidget(self.invoice_text)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        print_btn = QPushButton("Print Invoice")
        print_btn.clicked.connect(self.print_invoice)
        
        export_pdf_btn = QPushButton("Export PDF")
        export_pdf_btn.clicked.connect(lambda: self.export_invoice('pdf'))
        
        export_word_btn = QPushButton("Export Word")
        export_word_btn.clicked.connect(lambda: self.export_invoice('word'))
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        
        button_layout.addWidget(print_btn)
        button_layout.addWidget(export_pdf_btn)
        button_layout.addWidget(export_word_btn)
        button_layout.addStretch()
        button_layout.addWidget(close_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
        
        # Generate invoice HTML
        self.generate_invoice_html()
    
    def generate_invoice_html(self):
        # Get logo path
        logo_html = ""
        if os.path.exists(COMPANY_LOGO_PATH):
            logo_html = f'<div class="invoice-header"><img src="{COMPANY_LOGO_PATH}" class="logo" style="max-width: 150px;"></div>'
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .invoice-header {{ text-align: center; margin-bottom: 30px; }}
                .invoice-title {{ font-size: 24px; font-weight: bold; color: #4CAF50; }}
                .invoice-details {{ margin: 20px 0; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th {{ background-color: #4CAF50; color: white; padding: 12px; text-align: left; }}
                td {{ padding: 10px; border-bottom: 1px solid #ddd; }}
                .total {{ font-size: 18px; font-weight: bold; text-align: right; margin-top: 20px; color: #4CAF50; }}
                .footer {{ margin-top: 50px; text-align: center; font-size: 12px; color: #666; }}
                .logo {{ max-width: 150px; margin-bottom: 10px; }}
            </style>
        </head>
        <body>
        {logo_html}
        
        <div class="invoice-header">
            <div class="invoice-title">{self.parent.software_name if self.parent else 'Business Software'}</div>
            <div>{COMPANY_NAME}</div>
        </div>
        
        <div class="invoice-details">
            <table>
                <tr>
                    <td><strong>Invoice No:</strong> {self.invoice_data['invoice_no']}</td>
                    <td><strong>Date:</strong> {self.invoice_data['date']}</td>
                </tr>
                <tr>
                    <td><strong>Time:</strong> {self.invoice_data['time']}</td>
                    <td><strong>Customer:</strong> {self.invoice_data.get('customer', 'Walk-in Customer')}</td>
                </tr>
            </table>
        </div>
        
        <table>
            <thead>
                <tr>
                    <th>Product</th>
                    <th>Quantity</th>
                    <th>Purchase Price</th>
                    <th>Selling Price</th>
                    <th>Profit</th>
                </tr>
            </thead>
            <tbody>
        """
        
        total_profit = 0
        for item in self.invoice_data['items']:
            html += f"""
                    <tr>
                        <td>{item['name']}</td>
                        <td>{item['quantity']}</td>
                        <td>${item['purchase_price']:.2f}</td>
                        <td>${item['selling_price']:.2f}</td>
                        <td>${item['profit']:.2f}</td>
                    </tr>
            """
            total_profit += item['profit']
        
        html += f"""
                </tbody>
            </table>
            
            <div class="total">
                Total Profit: ${total_profit:.2f}
            </div>
            
            <div class="footer">
                <p>Generated by {self.parent.software_name if self.parent else 'Business Software'}</p>
                <p>Developed by {DEVELOPER_NAME}</p>
                <p>WhatsApp: {DEVELOPER_WHATSAPP}</p>
            </div>
        </body>
        </html>
        """
        
        self.invoice_text.setHtml(html)
    
    def print_invoice(self):
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        dialog = QPrintDialog(printer, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.invoice_text.print(printer)
    
    def export_invoice(self, format_type):
        if format_type == 'pdf':
            file_path, _ = QFileDialog.getSaveFileName(self, "Save PDF", "", "PDF Files (*.pdf)")
            if file_path:
                self.export_to_pdf(file_path)
        elif format_type == 'word':
            file_path, _ = QFileDialog.getSaveFileName(self, "Save Word Document", "", "Word Files (*.docx)")
            if file_path:
                self.export_to_word(file_path)
    
    def export_to_pdf(self, file_path):
        doc = SimpleDocTemplate(file_path, pagesize=A4)
        story = []
        
        styles = getSampleStyleSheet()
        
        # Add title
        title = Paragraph(f"Invoice #{self.invoice_data['invoice_no']}", styles['Title'])
        story.append(title)
        
        # Add date
        date_text = Paragraph(f"Date: {self.invoice_data['date']} {self.invoice_data['time']}", styles['Normal'])
        story.append(date_text)
        
        # Create table data
        data = [['Product', 'Quantity', 'Purchase Price', 'Selling Price', 'Profit']]
        for item in self.invoice_data['items']:
            data.append([
                item['name'],
                str(item['quantity']),
                f"${item['purchase_price']:.2f}",
                f"${item['selling_price']:.2f}",
                f"${item['profit']:.2f}"
            ])
        
        # Add total row
        total_profit = sum(item['profit'] for item in self.invoice_data['items'])
        data.append(['', '', '', 'Total Profit:', f"${total_profit:.2f}"])
        
        # Create table
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -2), colors.beige),
            ('GRID', (0, 0), (-1, -2), 1, colors.black)
        ]))
        
        story.append(table)
        
        # Build PDF
        doc.build(story)
        QMessageBox.information(self, "Success", "PDF exported successfully!")
    
    def export_to_word(self, file_path):
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Invoice #{self.invoice_data["invoice_no"]}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        doc.add_paragraph(f'Date: {self.invoice_data["date"]} {self.invoice_data["time"]}')
        
        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Add headers
        headers = ['Product', 'Quantity', 'Purchase Price', 'Selling Price', 'Profit']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for item in self.invoice_data['items']:
            row = table.add_row()
            row.cells[0].text = item['name']
            row.cells[1].text = str(item['quantity'])
            row.cells[2].text = f"${item['purchase_price']:.2f}"
            row.cells[3].text = f"${item['selling_price']:.2f}"
            row.cells[4].text = f"${item['profit']:.2f}"
        
        # Add total
        total_profit = sum(item['profit'] for item in self.invoice_data['items'])
        doc.add_paragraph(f'Total Profit: ${total_profit:.2f}')
        
        doc.save(file_path)
        QMessageBox.information(self, "Success", "Word document exported successfully!")