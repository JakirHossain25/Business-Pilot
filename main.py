# main.py
import sys
import os
from datetime import datetime
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QTabWidget, QPushButton, QLabel, QLineEdit, QTableWidget, 
                             QTableWidgetItem, QMessageBox, QFileDialog, QGroupBox, 
                             QGridLayout, QFormLayout, QSpinBox, QDoubleSpinBox, 
                             QComboBox, QDateEdit, QTextEdit, QSplitter, QHeaderView,
                             QProgressBar, QStackedWidget, QListWidget, QListWidgetItem,
                             QDialog, QDialogButtonBox, QToolBar, QStatusBar, QMenuBar,
                             QMenu, QFrame, QScrollArea, QSplashScreen)
from PyQt6.QtCore import Qt, QTimer, QDateTime, QDate, pyqtSignal, QSize
from PyQt6.QtGui import QIcon, QPixmap, QFont, QAction, QColor, QPalette, QLinearGradient
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog
import sqlite3
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
import shutil
from datetime import datetime, timedelta
import json

class LoginDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Login - Business Management Software")
        self.setModal(True)
        self.setFixedSize(400, 300)
        
        # Set window flags to make it always on top and frameless
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint)
        
        layout = QVBoxLayout()
        
        # Logo placeholder
        logo_label = QLabel()
        logo_pixmap = QPixmap(100, 100)
        logo_pixmap.fill(Qt.GlobalColor.gray)
        logo_label.setPixmap(logo_pixmap)
        logo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(logo_label)
        
        # Title
        title = QLabel("Business Management System")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("font-size: 18px; font-weight: bold; margin: 10px;")
        layout.addWidget(title)
        
        # Form
        form_layout = QFormLayout()
        
        self.username = QLineEdit()
        self.username.setPlaceholderText("Enter username")
        self.username.setStyleSheet("padding: 8px; font-size: 14px;")
        form_layout.addRow("Username:", self.username)
        
        self.password = QLineEdit()
        self.password.setPlaceholderText("Enter password")
        self.password.setEchoMode(QLineEdit.EchoMode.Password)
        self.password.setStyleSheet("padding: 8px; font-size: 14px;")
        form_layout.addRow("Password:", self.password)
        
        layout.addLayout(form_layout)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        login_btn = QPushButton("Login")
        login_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                font-size: 14px;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        login_btn.clicked.connect(self.accept)
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                padding: 10px;
                font-size: 14px;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #da190b;
            }
        """)
        cancel_btn.clicked.connect(self.reject)
        
        button_layout.addWidget(login_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        self.setLayout(layout)

class BrandingSettings(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.setWindowTitle("Branding Settings")
        self.setModal(True)
        self.setFixedSize(600, 500)
        
        layout = QVBoxLayout()
        
        # Software Name
        name_group = QGroupBox("Software Information")
        name_layout = QFormLayout()
        
        self.software_name = QLineEdit()
        self.software_name.setText(parent.software_name if parent else "Business Management Software")
        name_layout.addRow("Software Name:", self.software_name)
        
        self.developer_name = QLineEdit()
        self.developer_name.setText(parent.developer_name if parent else "Your Company Name")
        name_layout.addRow("Developer Name:", self.developer_name)
        
        self.developer_contact = QLineEdit()
        self.developer_contact.setText(parent.developer_contact if parent else "contact@example.com")
        name_layout.addRow("Developer Contact:", self.developer_contact)
        
        self.developer_info = QTextEdit()
        self.developer_info.setMaximumHeight(80)
        self.developer_info.setText(parent.developer_info if parent else "Developed by Your Company\nVersion 1.0")
        name_layout.addRow("Developer Info:", self.developer_info)
        
        name_group.setLayout(name_layout)
        layout.addWidget(name_group)
        
        # Logo Settings
        logo_group = QGroupBox("Logo Settings")
        logo_layout = QVBoxLayout()
        
        self.logo_path = QLineEdit()
        self.logo_path.setReadOnly(True)
        self.logo_path.setText(parent.logo_path if parent and hasattr(parent, 'logo_path') else "")
        
        logo_btn_layout = QHBoxLayout()
        browse_btn = QPushButton("Browse Logo")
        browse_btn.clicked.connect(self.browse_logo)
        remove_btn = QPushButton("Remove Logo")
        remove_btn.clicked.connect(self.remove_logo)
        
        logo_btn_layout.addWidget(browse_btn)
        logo_btn_layout.addWidget(remove_btn)
        
        logo_layout.addWidget(QLabel("Current Logo:"))
        logo_layout.addWidget(self.logo_path)
        logo_layout.addLayout(logo_btn_layout)
        
        # Preview
        self.logo_preview = QLabel()
        self.logo_preview.setFixedSize(200, 200)
        self.logo_preview.setStyleSheet("border: 1px solid #ccc;")
        self.logo_preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        logo_layout.addWidget(self.logo_preview)
        
        logo_group.setLayout(logo_layout)
        layout.addWidget(logo_group)
        
        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(self.save_settings)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
        self.setLayout(layout)
        
        # Load existing logo if available
        if parent and hasattr(parent, 'logo_path') and parent.logo_path:
            self.update_logo_preview(parent.logo_path)
    
    def browse_logo(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Logo", "", "Images (*.png *.jpg *.jpeg *.bmp)")
        if file_path:
            self.logo_path.setText(file_path)
            self.update_logo_preview(file_path)
    
    def remove_logo(self):
        self.logo_path.clear()
        self.logo_preview.clear()
        self.logo_preview.setText("No Logo")
    
    def update_logo_preview(self, path):
        pixmap = QPixmap(path)
        if not pixmap.isNull():
            scaled_pixmap = pixmap.scaled(200, 200, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            self.logo_preview.setPixmap(scaled_pixmap)
    
    def save_settings(self):
        self.parent.software_name = self.software_name.text()
        self.parent.developer_name = self.developer_name.text()
        self.parent.developer_contact = self.developer_contact.text()
        self.parent.developer_info = self.developer_info.toPlainText()
        self.parent.logo_path = self.logo_path.text()
        
        # Update window title
        self.parent.setWindowTitle(f"{self.parent.software_name} - {self.parent.developer_name}")
        
        # Update status bar
        self.parent.status_bar.showMessage(f"Welcome to {self.parent.software_name}")
        
        self.accept()

class AddProductDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add Product")
        self.setModal(True)
        self.setFixedSize(500, 400)
        
        layout = QVBoxLayout()
        
        form_layout = QFormLayout()
        
        self.product_name = QLineEdit()
        self.product_name.setPlaceholderText("Enter product name")
        form_layout.addRow("Product Name:", self.product_name)
        
        self.purchase_price = QDoubleSpinBox()
        self.purchase_price.setRange(0, 999999)
        self.purchase_price.setPrefix("$ ")
        form_layout.addRow("Purchase Price:", self.purchase_price)
        
        self.selling_price = QDoubleSpinBox()
        self.selling_price.setRange(0, 999999)
        self.selling_price.setPrefix("$ ")
        form_layout.addRow("Selling Price:", self.selling_price)
        
        self.quantity = QSpinBox()
        self.quantity.setRange(0, 999999)
        form_layout.addRow("Quantity:", self.quantity)
        
        self.profit_label = QLabel("$ 0.00")
        self.profit_label.setStyleSheet("font-weight: bold; color: #4CAF50;")
        form_layout.addRow("Total Profit:", self.profit_label)
        
        layout.addLayout(form_layout)
        
        # Connect signals for auto profit calculation
        self.purchase_price.valueChanged.connect(self.calculate_profit)
        self.selling_price.valueChanged.connect(self.calculate_profit)
        self.quantity.valueChanged.connect(self.calculate_profit)
        
        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
        self.setLayout(layout)
    
    def calculate_profit(self):
        purchase = self.purchase_price.value()
        selling = self.selling_price.value()
        quantity = self.quantity.value()
        profit = (selling - purchase) * quantity
        self.profit_label.setText(f"$ {profit:.2f}")
    
    def get_product_data(self):
        return {
            'name': self.product_name.text(),
            'purchase_price': self.purchase_price.value(),
            'selling_price': self.selling_price.value(),
            'quantity': self.quantity.value(),
            'profit': (self.selling_price.value() - self.purchase_price.value()) * self.quantity.value()
        }

class InvoiceDialog(QDialog):
    def __init__(self, parent=None, invoice_data=None):
        super().__init__(parent)
        self.parent = parent
        self.invoice_data = invoice_data
        self.setWindowTitle(f"Invoice #{invoice_data['invoice_no']}")
        self.setModal(True)
        self.setMinimumSize(800, 600)
        
        layout = QVBoxLayout()
        
        # Invoice display area
        self.invoice_text = QTextEdit()
        self.invoice_text.setReadOnly(True)
        layout.addWidget(self.invoice_text)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        print_btn = QPushButton("Print Invoice")
        print_btn.clicked.connect(self.print_invoice)
        
        export_pdf_btn = QPushButton("Export PDF")
        export_pdf_btn.clicked.connect(lambda: self.export_invoice('pdf'))
        
        export_word_btn = QPushButton("Export Word")
        export_word_btn.clicked.connect(lambda: self.export_invoice('word'))
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        
        button_layout.addWidget(print_btn)
        button_layout.addWidget(export_pdf_btn)
        button_layout.addWidget(export_word_btn)
        button_layout.addStretch()
        button_layout.addWidget(close_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
        
        # Generate invoice HTML
        self.generate_invoice_html()
    
    def generate_invoice_html(self):
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .invoice-header {{ text-align: center; margin-bottom: 30px; }}
                .invoice-title {{ font-size: 24px; font-weight: bold; color: #333; }}
                .invoice-details {{ margin: 20px 0; }}
                .company-info {{ margin-bottom: 20px; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th {{ background-color: #4CAF50; color: white; padding: 12px; text-align: left; }}
                td {{ padding: 10px; border-bottom: 1px solid #ddd; }}
                .total {{ font-size: 18px; font-weight: bold; text-align: right; margin-top: 20px; }}
                .footer {{ margin-top: 50px; text-align: center; font-size: 12px; color: #666; }}
                .logo {{ max-width: 150px; margin-bottom: 10px; }}
            </style>
        </head>
        <body>
        """
        
        # Add logo if available
        if self.parent and self.parent.logo_path and os.path.exists(self.parent.logo_path):
            html += f'<div class="invoice-header"><img src="{self.parent.logo_path}" class="logo"></div>'
        
        html += f"""
            <div class="invoice-header">
                <div class="invoice-title">{self.parent.software_name if self.parent else 'Business Software'}</div>
                <div>{self.parent.company_name if self.parent else 'Your Company'}</div>
            </div>
            
            <div class="invoice-details">
                <table>
                    <tr>
                        <td><strong>Invoice No:</strong> {self.invoice_data['invoice_no']}</td>
                        <td><strong>Date:</strong> {self.invoice_data['date']}</td>
                    </tr>
                    <tr>
                        <td><strong>Time:</strong> {self.invoice_data['time']}</td>
                        <td><strong>Customer:</strong> {self.invoice_data.get('customer', 'Walk-in Customer')}</td>
                    </tr>
                </table>
            </div>
            
            <table>
                <thead>
                    <tr>
                        <th>Product</th>
                        <th>Quantity</th>
                        <th>Purchase Price</th>
                        <th>Selling Price</th>
                        <th>Profit</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        total_profit = 0
        for item in self.invoice_data['items']:
            html += f"""
                    <tr>
                        <td>{item['name']}</td>
                        <td>{item['quantity']}</td>
                        <td>${item['purchase_price']:.2f}</td>
                        <td>${item['selling_price']:.2f}</td>
                        <td>${item['profit']:.2f}</td>
                    </tr>
            """
            total_profit += item['profit']
        
        html += f"""
                </tbody>
            </table>
            
            <div class="total">
                Total Profit: ${total_profit:.2f}
            </div>
            
            <div class="footer">
                <p>Generated by {self.parent.software_name if self.parent else 'Business Software'}</p>
                <p>Developed by {self.parent.developer_name if self.parent else 'Your Company'}</p>
                <p>{self.parent.developer_contact if self.parent else 'contact@example.com'}</p>
            </div>
        </body>
        </html>
        """
        
        self.invoice_text.setHtml(html)
    
    def print_invoice(self):
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        dialog = QPrintDialog(printer, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.invoice_text.print(printer)
    
    def export_invoice(self, format_type):
        if format_type == 'pdf':
            file_path, _ = QFileDialog.getSaveFileName(self, "Save PDF", "", "PDF Files (*.pdf)")
            if file_path:
                self.export_to_pdf(file_path)
        elif format_type == 'word':
            file_path, _ = QFileDialog.getSaveFileName(self, "Save Word Document", "", "Word Files (*.docx)")
            if file_path:
                self.export_to_word(file_path)
    
    def export_to_pdf(self, file_path):
        doc = SimpleDocTemplate(file_path, pagesize=A4)
        story = []
        
        styles = getSampleStyleSheet()
        
        # Add title
        title = Paragraph(f"Invoice #{self.invoice_data['invoice_no']}", styles['Title'])
        story.append(title)
        
        # Add date
        date_text = Paragraph(f"Date: {self.invoice_data['date']} {self.invoice_data['time']}", styles['Normal'])
        story.append(date_text)
        
        # Create table data
        data = [['Product', 'Quantity', 'Purchase Price', 'Selling Price', 'Profit']]
        for item in self.invoice_data['items']:
            data.append([
                item['name'],
                str(item['quantity']),
                f"${item['purchase_price']:.2f}",
                f"${item['selling_price']:.2f}",
                f"${item['profit']:.2f}"
            ])
        
        # Add total row
        total_profit = sum(item['profit'] for item in self.invoice_data['items'])
        data.append(['', '', '', 'Total Profit:', f"${total_profit:.2f}"])
        
        # Create table
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -2), colors.beige),
            ('GRID', (0, 0), (-1, -2), 1, colors.black)
        ]))
        
        story.append(table)
        
        # Build PDF
        doc.build(story)
        QMessageBox.information(self, "Success", "PDF exported successfully!")
    
    def export_to_word(self, file_path):
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Invoice #{self.invoice_data["invoice_no"]}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        doc.add_paragraph(f'Date: {self.invoice_data["date"]} {self.invoice_data["time"]}')
        
        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Add headers
        headers = ['Product', 'Quantity', 'Purchase Price', 'Selling Price', 'Profit']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for item in self.invoice_data['items']:
            row = table.add_row()
            row.cells[0].text = item['name']
            row.cells[1].text = str(item['quantity'])
            row.cells[2].text = f"${item['purchase_price']:.2f}"
            row.cells[3].text = f"${item['selling_price']:.2f}"
            row.cells[4].text = f"${item['profit']:.2f}"
        
        # Add total
        total_profit = sum(item['profit'] for item in self.invoice_data['items'])
        doc.add_paragraph(f'Total Profit: ${total_profit:.2f}')
        
        doc.save(file_path)
        QMessageBox.information(self, "Success", "Word document exported successfully!")

class ReportsWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.setup_ui()
    
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # Report controls
        control_layout = QHBoxLayout()
        
        self.report_type = QComboBox()
        self.report_type.addItems(["Weekly Report", "Monthly Report"])
        control_layout.addWidget(QLabel("Report Type:"))
        control_layout.addWidget(self.report_type)
        
        generate_btn = QPushButton("Generate Report")
        generate_btn.clicked.connect(self.generate_report)
        control_layout.addWidget(generate_btn)
        
        export_btn = QPushButton("Export Report")
        export_btn.clicked.connect(self.export_report)
        control_layout.addWidget(export_btn)
        
        print_btn = QPushButton("Print Report")
        print_btn.clicked.connect(self.print_report)
        control_layout.addWidget(print_btn)
        
        control_layout.addStretch()
        layout.addLayout(control_layout)
        
        # Chart area
        self.figure = Figure(figsize=(8, 6))
        self.canvas = FigureCanvas(self.figure)
        layout.addWidget(self.canvas)
        
        self.setLayout(layout)
        
        # Generate initial report
        self.generate_report()
    
    def generate_report(self):
        # Clear figure
        self.figure.clear()
        
        # Create subplot
        ax = self.figure.add_subplot(111)
        
        # Get data from database
        if self.report_type.currentText() == "Weekly Report":
            data = self.get_weekly_data()
            ax.set_title('Weekly Profit/Loss Report')
        else:
            data = self.get_monthly_data()
            ax.set_title('Monthly Profit/Loss Report')
        
        if data:
            # Unzip the data
            dates = [row[0] for row in data]
            profits = [row[1] for row in data]
            
            ax.bar(dates, profits)
            ax.set_xlabel('Date')
            ax.set_ylabel('Profit/Loss ($)')
            ax.tick_params(axis='x', rotation=45)
            
            # Color bars based on profit/loss
            for i, bar in enumerate(ax.patches):
                if profits[i] >= 0:
                    bar.set_color('green')
                else:
                    bar.set_color('red')
        else:
            ax.text(0.5, 0.5, 'No data available for this period', 
                   horizontalalignment='center', verticalalignment='center',
                   transform=ax.transAxes)
        
        self.figure.tight_layout()
        self.canvas.draw()
    
    def get_weekly_data(self):
        # Query database for weekly data using invoice date from invoices table
        cursor = self.parent.conn.cursor()
        end_date = datetime.now()
        start_date = end_date - timedelta(days=7)
        
        cursor.execute("""
            SELECT i.date, SUM(ii.profit) as daily_profit
            FROM invoices i
            LEFT JOIN invoice_items ii ON i.id = ii.invoice_id
            WHERE i.date >= ? AND i.date <= ?
            GROUP BY i.date
            ORDER BY i.date
        """, (start_date.strftime('%Y-%m-%d'), end_date.strftime('%Y-%m-%d')))
        
        return cursor.fetchall()
    
    def get_monthly_data(self):
        # Query database for monthly data
        cursor = self.parent.conn.cursor()
        end_date = datetime.now()
        start_date = end_date - timedelta(days=30)
        
        cursor.execute("""
            SELECT i.date, SUM(ii.profit) as daily_profit
            FROM invoices i
            LEFT JOIN invoice_items ii ON i.id = ii.invoice_id
            WHERE i.date >= ? AND i.date <= ?
            GROUP BY i.date
            ORDER BY i.date
        """, (start_date.strftime('%Y-%m-%d'), end_date.strftime('%Y-%m-%d')))
        
        return cursor.fetchall()
    
    def export_report(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Export Report", "", "Excel Files (*.xlsx);;PDF Files (*.pdf)")
        if file_path:
            if file_path.endswith('.xlsx'):
                self.export_to_excel(file_path)
            elif file_path.endswith('.pdf'):
                self.export_to_pdf(file_path)
    
    def export_to_excel(self, file_path):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Report"
        
        # Add headers
        headers = ['Date', 'Profit/Loss']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
            cell.fill = PatternFill(start_color="4CAF50", end_color="4CAF50", fill_type="solid")
        
        # Add data
        data = self.get_weekly_data() if self.report_type.currentText() == "Weekly Report" else self.get_monthly_data()
        for row, (date, profit) in enumerate(data, 2):
            ws.cell(row=row, column=1, value=date)
            ws.cell(row=row, column=2, value=profit if profit else 0)
        
        # Save file
        wb.save(file_path)
        QMessageBox.information(self, "Success", "Report exported successfully!")
    
    def export_to_pdf(self, file_path):
        # Save current figure to PDF
        self.figure.savefig(file_path, format='pdf', bbox_inches='tight')
        QMessageBox.information(self, "Success", "Report exported successfully!")
    
    def print_report(self):
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        dialog = QPrintDialog(printer, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            painter = QPainter(printer)
            self.figure.savefig(painter, format='pdf')
            painter.end()

class DashboardWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.setup_ui()
        
        # Setup timer for updates
        self.timer = QTimer()
        self.timer.timeout.connect(self.update_dashboard)
        self.timer.start(60000)  # Update every minute
    
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # Welcome message
        welcome_label = QLabel(f"Welcome to {self.parent.software_name if self.parent else 'Business Software'}")
        welcome_label.setStyleSheet("font-size: 24px; font-weight: bold; margin: 10px;")
        welcome_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(welcome_label)
        
        # Stats cards
        stats_layout = QHBoxLayout()
        
        # Sales card
        sales_card = self.create_stat_card("Total Sales", "$0", "#4CAF50")
        self.sales_label = sales_card.findChild(QLabel, "value_label")
        stats_layout.addWidget(sales_card)
        
        # Profit card
        profit_card = self.create_stat_card("Total Profit", "$0", "#2196F3")
        self.profit_label = profit_card.findChild(QLabel, "value_label")
        stats_layout.addWidget(profit_card)
        
        # Products card
        products_card = self.create_stat_card("Total Products", "0", "#FF9800")
        self.products_label = products_card.findChild(QLabel, "value_label")
        stats_layout.addWidget(products_card)
        
        # Invoices card
        invoices_card = self.create_stat_card("Total Invoices", "0", "#9C27B0")
        self.invoices_label = invoices_card.findChild(QLabel, "value_label")
        stats_layout.addWidget(invoices_card)
        
        layout.addLayout(stats_layout)
        
        # Progress wheels
        progress_layout = QHBoxLayout()
        
        # Sales progress
        sales_progress_group = QGroupBox("Sales Progress")
        sales_progress_layout = QVBoxLayout()
        self.sales_progress = QProgressBar()
        self.sales_progress.setRange(0, 100)
        self.sales_progress.setValue(0)
        self.sales_progress.setFormat("%p%")
        self.sales_progress.setStyleSheet("""
            QProgressBar {
                border: 2px solid grey;
                border-radius: 5px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #4CAF50;
                width: 10px;
            }
        """)
        sales_progress_layout.addWidget(self.sales_progress)
        sales_progress_group.setLayout(sales_progress_layout)
        progress_layout.addWidget(sales_progress_group)
        
        # Profit progress
        profit_progress_group = QGroupBox("Profit Progress")
        profit_progress_layout = QVBoxLayout()
        self.profit_progress = QProgressBar()
        self.profit_progress.setRange(0, 100)
        self.profit_progress.setValue(0)
        self.profit_progress.setFormat("%p%")
        self.profit_progress.setStyleSheet("""
            QProgressBar {
                border: 2px solid grey;
                border-radius: 5px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #2196F3;
                width: 10px;
            }
        """)
        profit_progress_layout.addWidget(self.profit_progress)
        profit_progress_group.setLayout(profit_progress_layout)
        progress_layout.addWidget(profit_progress_group)
        
        # Time tracking
        time_progress_group = QGroupBox("Time Tracking")
        time_progress_layout = QVBoxLayout()
        self.time_label = QLabel()
        self.time_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.time_label.setStyleSheet("font-size: 18px; font-weight: bold;")
        time_progress_layout.addWidget(self.time_label)
        time_progress_group.setLayout(time_progress_layout)
        progress_layout.addWidget(time_progress_group)
        
        layout.addLayout(progress_layout)
        
        # Recent activity
        recent_group = QGroupBox("Recent Activity")
        recent_layout = QVBoxLayout()
        self.recent_table = QTableWidget()
        self.recent_table.setColumnCount(4)
        self.recent_table.setHorizontalHeaderLabels(["Invoice No", "Date", "Items", "Profit"])
        self.recent_table.horizontalHeader().setStretchLastSection(True)
        recent_layout.addWidget(self.recent_table)
        recent_group.setLayout(recent_layout)
        layout.addWidget(recent_group)
        
        self.setLayout(layout)
        
        # Initial update
        self.update_dashboard()
    
    def create_stat_card(self, title, value, color):
        card = QFrame()
        card.setFrameStyle(QFrame.Shape.Box)
        card.setStyleSheet(f"""
            QFrame {{
                background-color: white;
                border: 2px solid {color};
                border-radius: 10px;
                padding: 10px;
                margin: 5px;
            }}
        """)
        
        layout = QVBoxLayout()
        
        title_label = QLabel(title)
        title_label.setStyleSheet("font-size: 14px; color: #666;")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)
        
        value_label = QLabel(value)
        value_label.setObjectName("value_label")
        value_label.setStyleSheet(f"font-size: 24px; font-weight: bold; color: {color};")
        value_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(value_label)
        
        card.setLayout(layout)
        return card
    
    def update_dashboard(self):
        cursor = self.parent.conn.cursor()
        
        # Update sales
        cursor.execute("SELECT SUM(selling_price * quantity) FROM invoice_items")
        total_sales = cursor.fetchone()[0] or 0
        self.sales_label.setText(f"${total_sales:.2f}")
        
        # Update profit
        cursor.execute("SELECT SUM(profit) FROM invoice_items")
        total_profit = cursor.fetchone()[0] or 0
        self.profit_label.setText(f"${total_profit:.2f}")
        
        # Update products count
        cursor.execute("SELECT COUNT(*) FROM products")
        total_products = cursor.fetchone()[0] or 0
        self.products_label.setText(str(total_products))
        
        # Update invoices count
        cursor.execute("SELECT COUNT(*) FROM invoices")
        total_invoices = cursor.fetchone()[0] or 0
        self.invoices_label.setText(str(total_invoices))
        
        # Update progress bars
        monthly_target = 10000  # Example target
        if total_sales > 0:
            self.sales_progress.setValue(min(int((total_sales / monthly_target) * 100), 100))
        if total_profit > 0:
            self.profit_progress.setValue(min(int((total_profit / (monthly_target * 0.2)) * 100), 100))
        
        # Update time
        current_time = QDateTime.currentDateTime()
        self.time_label.setText(current_time.toString("hh:mm AP"))
        
        # Update recent activity
        cursor.execute("""
            SELECT i.invoice_no, i.date, COUNT(ii.id) as items, SUM(ii.profit) as profit
            FROM invoices i
            LEFT JOIN invoice_items ii ON i.id = ii.invoice_id
            GROUP BY i.id
            ORDER BY i.date DESC, i.time DESC
            LIMIT 10
        """)
        recent_data = cursor.fetchall()
        
        self.recent_table.setRowCount(len(recent_data))
        for row, data in enumerate(recent_data):
            for col, value in enumerate(data):
                if col == 3 and value:  # Profit column
                    self.recent_table.setItem(row, col, QTableWidgetItem(f"${value:.2f}"))
                else:
                    self.recent_table.setItem(row, col, QTableWidgetItem(str(value)))

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        
        # Initialize branding
        self.software_name = "Business Management Software"
        self.developer_name = "Your Company Name"
        self.developer_contact = "contact@example.com"
        self.developer_info = "Developed by Your Company\nVersion 1.0"
        self.company_name = "Your Business Name"
        self.logo_path = ""
        
        # Initialize database
        self.init_database()
        
        # Setup UI
        self.setup_ui()
        
        # Load settings
        self.load_settings()
        
        # Show login dialog
        self.show_login()
    
    def init_database(self):
        self.conn = sqlite3.connect('business_management.db')
        cursor = self.conn.cursor()
        
        # Create tables
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS products (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                purchase_price REAL NOT NULL,
                selling_price REAL NOT NULL,
                quantity INTEGER NOT NULL,
                profit REAL NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS invoices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_no TEXT UNIQUE NOT NULL,
                date TEXT NOT NULL,
                time TEXT NOT NULL,
                customer TEXT,
                total_profit REAL NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS invoice_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_id INTEGER NOT NULL,
                product_id INTEGER NOT NULL,
                name TEXT NOT NULL,
                quantity INTEGER NOT NULL,
                purchase_price REAL NOT NULL,
                selling_price REAL NOT NULL,
                profit REAL NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (invoice_id) REFERENCES invoices (id),
                FOREIGN KEY (product_id) REFERENCES products (id)
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS settings (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        ''')
        
        self.conn.commit()
    
    def setup_ui(self):
        self.setWindowTitle(f"{self.software_name} - {self.developer_name}")
        self.setGeometry(100, 100, 1400, 800)
        
        # Set application icon and style
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                font-size: 14px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:pressed {
                background-color: #3d8b40;
            }
            QTabWidget::pane {
                border: 1px solid #ccc;
                background-color: white;
            }
            QTabBar::tab {
                background-color: #e0e0e0;
                padding: 8px 16px;
                margin-right: 2px;
            }
            QTabBar::tab:selected {
                background-color: white;
            }
            QTableWidget {
                gridline-color: #ddd;
            }
            QTableWidget::item {
                padding: 5px;
            }
            QHeaderView::section {
                background-color: #4CAF50;
                color: white;
                padding: 8px;
                border: none;
            }
        """)
        
        # Create menu bar
        self.create_menu_bar()
        
        # Create toolbar
        self.create_toolbar()
        
        # Create central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Main layout
        main_layout = QVBoxLayout(central_widget)
        
        # Create tab widget
        self.tab_widget = QTabWidget()
        
        # Dashboard tab
        self.dashboard = DashboardWidget(self)
        self.tab_widget.addTab(self.dashboard, "Dashboard")
        
        # Products tab
        self.products_tab = self.create_products_tab()
        self.tab_widget.addTab(self.products_tab, "Products")
        
        # Invoices tab
        self.invoices_tab = self.create_invoices_tab()
        self.tab_widget.addTab(self.invoices_tab, "Invoices")
        
        # Reports tab
        self.reports_tab = ReportsWidget(self)
        self.tab_widget.addTab(self.reports_tab, "Reports")
        
        main_layout.addWidget(self.tab_widget)
        
        # Create status bar
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage(f"Welcome to {self.software_name}")
        
        # Add developer info to status bar
        self.dev_label = QLabel(f"Developed by: {self.developer_name}")
        self.status_bar.addPermanentWidget(self.dev_label)
    
    def create_menu_bar(self):
        menubar = self.menuBar()
        
        # File menu
        file_menu = menubar.addMenu("File")
        
        # Settings action
        settings_action = QAction("Branding Settings", self)
        settings_action.triggered.connect(self.open_branding_settings)
        file_menu.addAction(settings_action)
        
        file_menu.addSeparator()
        
        # Backup action
        backup_action = QAction("Backup Database", self)
        backup_action.triggered.connect(self.backup_database)
        file_menu.addAction(backup_action)
        
        # Restore action
        restore_action = QAction("Restore Database", self)
        restore_action.triggered.connect(self.restore_database)
        file_menu.addAction(restore_action)
        
        file_menu.addSeparator()
        
        # Exit action
        exit_action = QAction("Exit", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Help menu
        help_menu = menubar.addMenu("Help")
        
        # About action
        about_action = QAction("About", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
    
    def create_toolbar(self):
        toolbar = QToolBar()
        self.addToolBar(toolbar)
        
        # Add product action
        add_product_action = QAction("Add Product", self)
        add_product_action.triggered.connect(self.add_product)
        toolbar.addAction(add_product_action)
        
        # Create invoice action
        create_invoice_action = QAction("Create Invoice", self)
        create_invoice_action.triggered.connect(self.generate_invoice)
        toolbar.addAction(create_invoice_action)
        
        toolbar.addSeparator()
        
        # Refresh action
        refresh_action = QAction("Refresh", self)
        refresh_action.triggered.connect(self.refresh_data)
        toolbar.addAction(refresh_action)
    
    def create_products_tab(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Search bar
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel("Search:"))
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search products...")
        self.search_input.textChanged.connect(self.search_products)
        search_layout.addWidget(self.search_input)
        
        search_btn = QPushButton("Search")
        search_btn.clicked.connect(self.search_products)
        search_layout.addWidget(search_btn)
        
        layout.addLayout(search_layout)
        
        # Products table
        self.products_table = QTableWidget()
        self.products_table.setColumnCount(6)
        self.products_table.setHorizontalHeaderLabels([
            "ID", "Product Name", "Purchase Price", "Selling Price", "Quantity", "Profit"
        ])
        self.products_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.products_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        layout.addWidget(self.products_table)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        add_btn = QPushButton("Add Product")
        add_btn.clicked.connect(self.add_product)
        button_layout.addWidget(add_btn)
        
        edit_btn = QPushButton("Edit Product")
        edit_btn.clicked.connect(self.edit_product)
        button_layout.addWidget(edit_btn)
        
        delete_btn = QPushButton("Delete Product")
        delete_btn.clicked.connect(self.delete_product)
        button_layout.addWidget(delete_btn)
        
        button_layout.addStretch()
        
        # Stock alert button
        stock_alert_btn = QPushButton("Stock Alert")
        stock_alert_btn.clicked.connect(self.check_stock_alert)
        stock_alert_btn.setStyleSheet("background-color: #FF9800;")
        button_layout.addWidget(stock_alert_btn)
        
        layout.addLayout(button_layout)
        
        # Load products
        self.load_products()
        
        return widget
    
    def create_invoices_tab(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Invoice creation area
        create_group = QGroupBox("Create Invoice")
        create_layout = QVBoxLayout()
        
        # Product selection
        product_layout = QHBoxLayout()
        product_layout.addWidget(QLabel("Select Product:"))
        self.invoice_product = QComboBox()
        self.load_products_to_combo()
        product_layout.addWidget(self.invoice_product)
        
        self.invoice_quantity = QSpinBox()
        self.invoice_quantity.setRange(1, 999)
        self.invoice_quantity.setValue(1)
        product_layout.addWidget(QLabel("Quantity:"))
        product_layout.addWidget(self.invoice_quantity)
        
        add_to_invoice_btn = QPushButton("Add to Invoice")
        add_to_invoice_btn.clicked.connect(self.add_to_invoice)
        product_layout.addWidget(add_to_invoice_btn)
        
        product_layout.addStretch()
        create_layout.addLayout(product_layout)
        
        # Invoice items table
        self.invoice_items_table = QTableWidget()
        self.invoice_items_table.setColumnCount(5)
        self.invoice_items_table.setHorizontalHeaderLabels([
            "Product", "Quantity", "Purchase Price", "Selling Price", "Profit"
        ])
        self.invoice_items_table.horizontalHeader().setStretchLastSection(True)
        create_layout.addWidget(self.invoice_items_table)
        
        # Remove item button
        remove_item_btn = QPushButton("Remove Selected Item")
        remove_item_btn.clicked.connect(self.remove_invoice_item)
        remove_item_btn.setStyleSheet("background-color: #f44336;")
        create_layout.addWidget(remove_item_btn)
        
        # Total
        self.invoice_total = QLabel("Total Profit: $0.00")
        self.invoice_total.setStyleSheet("font-size: 16px; font-weight: bold; color: #4CAF50;")
        self.invoice_total.setAlignment(Qt.AlignmentFlag.AlignRight)
        create_layout.addWidget(self.invoice_total)
        
        # Generate button
        generate_btn = QPushButton("Generate Invoice")
        generate_btn.clicked.connect(self.generate_invoice)
        generate_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                font-size: 16px;
                padding: 10px;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        create_layout.addWidget(generate_btn)
        
        create_group.setLayout(create_layout)
        layout.addWidget(create_group)
        
        # Recent invoices
        recent_group = QGroupBox("Recent Invoices")
        recent_layout = QVBoxLayout()
        
        self.invoices_table = QTableWidget()
        self.invoices_table.setColumnCount(5)
        self.invoices_table.setHorizontalHeaderLabels([
            "Invoice No", "Date", "Time", "Items", "Total Profit"
        ])
        self.invoices_table.horizontalHeader().setStretchLastSection(True)
        self.invoices_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.invoices_table.doubleClicked.connect(self.view_invoice_details)
        recent_layout.addWidget(self.invoices_table)
        
        recent_group.setLayout(recent_layout)
        layout.addWidget(recent_group)
        
        # Load invoices
        self.load_invoices()
        
        return widget
    
    def load_products(self):
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM products ORDER BY id DESC")
        products = cursor.fetchall()
        
        self.products_table.setRowCount(len(products))
        for row, product in enumerate(products):
            for col, value in enumerate(product):  # Include all columns
                if col in [2, 3, 5]:  # Price columns (purchase_price, selling_price, profit)
                    item = QTableWidgetItem(f"${value:.2f}")
                else:
                    item = QTableWidgetItem(str(value))
                self.products_table.setItem(row, col, item)
    
    def load_products_to_combo(self):
        self.invoice_product.clear()
        cursor = self.conn.cursor()
        cursor.execute("SELECT id, name FROM products WHERE quantity > 0")
        products = cursor.fetchall()
        for product in products:
            self.invoice_product.addItem(product[1], product[0])
    
    def load_invoices(self):
        cursor = self.conn.cursor()
        cursor.execute("""
            SELECT i.invoice_no, i.date, i.time, COUNT(ii.id) as items, i.total_profit
            FROM invoices i
            LEFT JOIN invoice_items ii ON i.id = ii.invoice_id
            GROUP BY i.id
            ORDER BY i.created_at DESC
            LIMIT 20
        """)
        invoices = cursor.fetchall()
        
        self.invoices_table.setRowCount(len(invoices))
        for row, invoice in enumerate(invoices):
            for col, value in enumerate(invoice):
                if col == 4:  # Profit column
                    item = QTableWidgetItem(f"${value:.2f}")
                else:
                    item = QTableWidgetItem(str(value))
                self.invoices_table.setItem(row, col, item)
    
    def add_product(self):
        dialog = AddProductDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            data = dialog.get_product_data()
            
            cursor = self.conn.cursor()
            cursor.execute('''
                INSERT INTO products (name, purchase_price, selling_price, quantity, profit)
                VALUES (?, ?, ?, ?, ?)
            ''', (data['name'], data['purchase_price'], data['selling_price'], 
                  data['quantity'], data['profit']))
            self.conn.commit()
            
            self.load_products()
            self.load_products_to_combo()
            QMessageBox.information(self, "Success", "Product added successfully!")
    
    def edit_product(self):
        current_row = self.products_table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "Warning", "Please select a product to edit.")
            return
        
        product_id = self.products_table.item(current_row, 0).text()
        
        # Get product data
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM products WHERE id = ?", (product_id,))
        product = cursor.fetchone()
        
        if product:
            dialog = AddProductDialog(self)
            dialog.setWindowTitle("Edit Product")
            dialog.product_name.setText(product[1])
            dialog.purchase_price.setValue(product[2])
            dialog.selling_price.setValue(product[3])
            dialog.quantity.setValue(product[4])
            
            if dialog.exec() == QDialog.DialogCode.Accepted:
                data = dialog.get_product_data()
                
                cursor.execute('''
                    UPDATE products 
                    SET name=?, purchase_price=?, selling_price=?, quantity=?, profit=?
                    WHERE id=?
                ''', (data['name'], data['purchase_price'], data['selling_price'],
                      data['quantity'], data['profit'], product_id))
                self.conn.commit()
                
                self.load_products()
                self.load_products_to_combo()
                QMessageBox.information(self, "Success", "Product updated successfully!")
    
    def delete_product(self):
        current_row = self.products_table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "Warning", "Please select a product to delete.")
            return
        
        reply = QMessageBox.question(self, "Confirm Delete", 
                                     "Are you sure you want to delete this product?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            product_id = self.products_table.item(current_row, 0).text()
            
            cursor = self.conn.cursor()
            cursor.execute("DELETE FROM products WHERE id = ?", (product_id,))
            self.conn.commit()
            
            self.load_products()
            self.load_products_to_combo()
            QMessageBox.information(self, "Success", "Product deleted successfully!")
    
    def search_products(self):
        search_text = self.search_input.text().lower()
        
        for row in range(self.products_table.rowCount()):
            product_name = self.products_table.item(row, 1).text().lower()
            if search_text in product_name:
                self.products_table.setRowHidden(row, False)
            else:
                self.products_table.setRowHidden(row, True)
    
    def check_stock_alert(self):
        cursor = self.conn.cursor()
        cursor.execute("SELECT name, quantity FROM products WHERE quantity < 10 ORDER BY quantity")
        low_stock = cursor.fetchall()
        
        if low_stock:
            alert_message = "Low Stock Alert!\n\n"
            for product in low_stock:
                alert_message += f"{product[0]}: {product[1]} units remaining\n"
            QMessageBox.warning(self, "Stock Alert", alert_message)
        else:
            QMessageBox.information(self, "Stock Alert", "All products have sufficient stock.")
    
    def add_to_invoice(self):
        if self.invoice_product.count() == 0:
            QMessageBox.warning(self, "Warning", "No products available with stock!")
            return
            
        product_id = self.invoice_product.currentData()
        quantity = self.invoice_quantity.value()
        
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM products WHERE id = ?", (product_id,))
        product = cursor.fetchone()
        
        if product and product[4] >= quantity:
            profit = (product[3] - product[2]) * quantity
            
            # Add to invoice table
            row = self.invoice_items_table.rowCount()
            self.invoice_items_table.insertRow(row)
            self.invoice_items_table.setItem(row, 0, QTableWidgetItem(product[1]))
            self.invoice_items_table.setItem(row, 1, QTableWidgetItem(str(quantity)))
            self.invoice_items_table.setItem(row, 2, QTableWidgetItem(f"${product[2]:.2f}"))
            self.invoice_items_table.setItem(row, 3, QTableWidgetItem(f"${product[3]:.2f}"))
            self.invoice_items_table.setItem(row, 4, QTableWidgetItem(f"${profit:.2f}"))
            
            # Store product data in item data
            item = self.invoice_items_table.item(row, 0)
            item.setData(Qt.ItemDataRole.UserRole, {
                'id': product_id,
                'name': product[1],
                'purchase_price': product[2],
                'selling_price': product[3],
                'quantity': quantity,
                'profit': profit
            })
            
            # Update total
            self.update_invoice_total()
        else:
            QMessageBox.warning(self, "Warning", "Insufficient stock!")
    
    def remove_invoice_item(self):
        current_row = self.invoice_items_table.currentRow()
        if current_row >= 0:
            self.invoice_items_table.removeRow(current_row)
            self.update_invoice_total()
    
    def update_invoice_total(self):
        total = 0
        for row in range(self.invoice_items_table.rowCount()):
            profit_item = self.invoice_items_table.item(row, 4)
            total += float(profit_item.text().replace('$', ''))
        
        self.invoice_total.setText(f"Total Profit: ${total:.2f}")
    
    def generate_invoice(self):
        if self.invoice_items_table.rowCount() == 0:
            QMessageBox.warning(self, "Warning", "No items in invoice!")
            return
        
        # Generate invoice number
        invoice_no = f"INV-{datetime.now().strftime('%Y%m%d%H%M%S')}"
        current_date = datetime.now().strftime('%Y-%m-%d')
        current_time = datetime.now().strftime('%H:%M:%S')
        
        # Calculate total profit
        total_profit = 0
        items = []
        for row in range(self.invoice_items_table.rowCount()):
            item_data = self.invoice_items_table.item(row, 0).data(Qt.ItemDataRole.UserRole)
            items.append(item_data)
            total_profit += item_data['profit']
        
        cursor = self.conn.cursor()
        
        # Create invoice
        cursor.execute('''
            INSERT INTO invoices (invoice_no, date, time, total_profit)
            VALUES (?, ?, ?, ?)
        ''', (invoice_no, current_date, current_time, total_profit))
        
        invoice_id = cursor.lastrowid
        
        # Add invoice items and update stock
        for item in items:
            cursor.execute('''
                INSERT INTO invoice_items (invoice_id, product_id, name, quantity, 
                                         purchase_price, selling_price, profit)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (invoice_id, item['id'], item['name'], item['quantity'],
                  item['purchase_price'], item['selling_price'], item['profit']))
            
            # Update product quantity
            cursor.execute('''
                UPDATE products 
                SET quantity = quantity - ? 
                WHERE id = ?
            ''', (item['quantity'], item['id']))
        
        self.conn.commit()
        
        # Show invoice
        invoice_data = {
            'invoice_no': invoice_no,
            'date': current_date,
            'time': current_time,
            'items': items
        }
        
        dialog = InvoiceDialog(self, invoice_data)
        dialog.exec()
        
        # Clear invoice items
        self.invoice_items_table.setRowCount(0)
        self.invoice_total.setText("Total Profit: $0.00")
        
        # Refresh data
        self.load_products()
        self.load_products_to_combo()
        self.load_invoices()
        self.dashboard.update_dashboard()
        
        QMessageBox.information(self, "Success", f"Invoice {invoice_no} generated successfully!")
    
    def view_invoice_details(self):
        current_row = self.invoices_table.currentRow()
        if current_row < 0:
            return
        
        invoice_no = self.invoices_table.item(current_row, 0).text()
        
        # Get invoice details
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM invoices WHERE invoice_no = ?", (invoice_no,))
        invoice = cursor.fetchone()
        
        if invoice:
            cursor.execute("""
                SELECT name, quantity, purchase_price, selling_price, profit
                FROM invoice_items 
                WHERE invoice_id = ?
            """, (invoice[0],))
            items = cursor.fetchall()
            
            invoice_items = []
            for item in items:
                invoice_items.append({
                    'name': item[0],
                    'quantity': item[1],
                    'purchase_price': item[2],
                    'selling_price': item[3],
                    'profit': item[4]
                })
            
            invoice_data = {
                'invoice_no': invoice[1],
                'date': invoice[2],
                'time': invoice[3],
                'items': invoice_items
            }
            
            dialog = InvoiceDialog(self, invoice_data)
            dialog.exec()
    
    def open_branding_settings(self):
        dialog = BrandingSettings(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.save_settings()
            self.status_bar.showMessage("Branding settings updated successfully!")
    
    def backup_database(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Backup Database", "", "Database Files (*.db)")
        if file_path:
            shutil.copy2('business_management.db', file_path)
            QMessageBox.information(self, "Success", "Database backed up successfully!")
    
    def restore_database(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Restore Database", "", "Database Files (*.db)")
        if file_path:
            reply = QMessageBox.question(self, "Confirm Restore", 
                                        "This will overwrite current data. Continue?",
                                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                shutil.copy2(file_path, 'business_management.db')
                self.conn.close()
                self.init_database()
                self.load_products()
                self.load_invoices()
                self.dashboard.update_dashboard()
                QMessageBox.information(self, "Success", "Database restored successfully!")
    
    def show_about(self):
        about_text = f"""
        <h2>{self.software_name}</h2>
        <p>Version 1.0</p>
        <p>Developed by: {self.developer_name}</p>
        <p>Contact: {self.developer_contact}</p>
        <br>
        <p>{self.developer_info}</p>
        """
        
        QMessageBox.about(self, "About", about_text)
    
    def refresh_data(self):
        self.load_products()
        self.load_products_to_combo()
        self.load_invoices()
        self.dashboard.update_dashboard()
        self.status_bar.showMessage("Data refreshed successfully!", 3000)
    
    def save_settings(self):
        settings = {
            'software_name': self.software_name,
            'developer_name': self.developer_name,
            'developer_contact': self.developer_contact,
            'developer_info': self.developer_info,
            'company_name': self.company_name,
            'logo_path': self.logo_path
        }
        
        cursor = self.conn.cursor()
        for key, value in settings.items():
            cursor.execute('''
                INSERT OR REPLACE INTO settings (key, value)
                VALUES (?, ?)
            ''', (key, str(value)))
        
        self.conn.commit()
    
    def load_settings(self):
        cursor = self.conn.cursor()
        cursor.execute("SELECT key, value FROM settings")
        settings = cursor.fetchall()
        
        for key, value in settings:
            if key == 'software_name':
                self.software_name = value
            elif key == 'developer_name':
                self.developer_name = value
            elif key == 'developer_contact':
                self.developer_contact = value
            elif key == 'developer_info':
                self.developer_info = value
            elif key == 'company_name':
                self.company_name = value
            elif key == 'logo_path':
                self.logo_path = value
    
    def show_login(self):
        dialog = LoginDialog()
        if dialog.exec() != QDialog.DialogCode.Accepted:
            sys.exit()

def main():
    app = QApplication(sys.argv)
    
    # Set application style
    app.setStyle('Fusion')
    
    # Create splash screen
    splash_pixmap = QPixmap(400, 300)
    splash_pixmap.fill(Qt.GlobalColor.white)
    splash = QSplashScreen(splash_pixmap)
    splash.show()
    splash.showMessage("Loading Business Management Software...", Qt.AlignmentFlag.AlignCenter, Qt.GlobalColor.black)
    
    # Process events to show splash
    app.processEvents()
    
    # Create main window
    window = MainWindow()
    
    # Close splash and show main window
    splash.finish(window)
    window.show()
    
    sys.exit(app.exec())

if __name__ == '__main__':
    main()